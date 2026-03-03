import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict
from backend.app.models import Questionnaire, Answer

def export_to_docx(questionnaire: Questionnaire, answers: List[Answer], output_path: str):
    """Export questionnaire with answers to DOCX format."""
    doc = Document()
    
    # Title
    title = doc.add_heading(f"Questionnaire: {questionnaire.filename}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    doc.add_paragraph(f"Generated: {questionnaire.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()
    
    # Questions and Answers
    for i, answer in enumerate(answers, 1):
        # Question
        q_heading = doc.add_heading(f"Question {i}", level=2)
        q_para = doc.add_paragraph(answer.question)
        q_para.runs[0].bold = True
        
        # Answer
        doc.add_heading("Answer", level=3)
        if answer.is_not_found:
            answer_text = "Not found in references."
            answer_para = doc.add_paragraph(answer_text)
            answer_para.runs[0].italic = True
            answer_para.runs[0].font.color.rgb = RGBColor(255, 0, 0)
        else:
            answer_para = doc.add_paragraph(answer.answer or "No answer generated.")
            
            # Citations
            if answer.citations:
                citations = json.loads(answer.citations)
                doc.add_heading("Citations", level=4)
                for citation in citations:
                    cite_text = f"- {citation.get('document', 'Unknown')}: {citation.get('position', 'N/A')}"
                    doc.add_paragraph(cite_text)
            
            # Evidence Snippets
            if answer.evidence_snippets:
                snippets = json.loads(answer.evidence_snippets)
                doc.add_heading("Evidence Snippets", level=4)
                for snippet in snippets:
                    snippet_para = doc.add_paragraph(snippet)
                    snippet_para.runs[0].italic = True
                    snippet_para.runs[0].font.size = Pt(9)
        
        # Confidence Score
        if answer.confidence_score:
            conf_para = doc.add_paragraph(f"Confidence Score: {answer.confidence_score:.2f}")
            conf_para.runs[0].font.size = Pt(10)
            conf_para.runs[0].font.color.rgb = RGBColor(0, 128, 0)
        
        doc.add_paragraph()  # Spacing
        doc.add_paragraph("-" * 60)  # Divider
        doc.add_paragraph()
    
    # Save
    doc.save(output_path)

def generate_coverage_summary(answers: List[Answer]) -> Dict:
    """Generate coverage summary statistics."""
    total = len(answers)
    answered = sum(1 for a in answers if a.answer and not a.is_not_found)
    not_found = sum(1 for a in answers if a.is_not_found)
    edited = sum(1 for a in answers if a.status == 'edited')
    
    avg_confidence = 0.0
    answered_with_scores = [a.confidence_score for a in answers if a.confidence_score]
    if answered_with_scores:
        avg_confidence = sum(answered_with_scores) / len(answered_with_scores)
    
    return {
        'total_questions': total,
        'answered': answered,
        'not_found': not_found,
        'edited': edited,
        'average_confidence': round(avg_confidence, 3)
    }
